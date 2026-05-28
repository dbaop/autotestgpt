import io
import tempfile
from pathlib import Path

from flask import Flask
from docx import Document
import werkzeug
from openpyxl import Workbook

if not hasattr(werkzeug, "__version__"):
    werkzeug.__version__ = "3"


def _build_test_app(tmp_dir: Path):
    from models import db
    from api import api_blueprint

    app = Flask(__name__)
    app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{tmp_dir / 'phase6_test.db'}"
    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
    db.init_app(app)
    app.register_blueprint(api_blueprint, url_prefix="/api")

    with app.app_context():
        db.drop_all()
        db.create_all()

    return app, db


def _local_tmp_dir() -> Path:
    workspace_tmp = Path("workspace") / "pytest_phase6"
    workspace_tmp.mkdir(parents=True, exist_ok=True)
    return Path(tempfile.mkdtemp(dir=workspace_tmp))


def _build_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Login"
    sheet["A1"] = "Feature"
    sheet["B1"] = "Rule"
    sheet["A2"] = "SMS Login"
    sheet["B2"] = "Retry limit after wrong codes"
    buffer = io.BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Login Requirement", level=1)
    document.add_paragraph("Users can sign in with SMS verification.")
    document.add_paragraph("Expired verification code should show a warning.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Scenario"
    table.cell(0, 1).text = "Expected"
    table.cell(1, 0).text = "Wrong code"
    table.cell(1, 1).text = "Block after retry limit"
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def test_requirement_import_supports_xlsx_files():
    app, _db = _build_test_app(_local_tmp_dir())
    client = app.test_client()

    response = client.post(
        "/api/requirements/import",
        data={
            "title": "Imported xlsx requirement",
            "file": (io.BytesIO(_build_xlsx_bytes()), "login.xlsx"),
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 201
    payload = response.get_json()
    assert "SMS Login" in payload["requirement"]["description"]


def test_requirement_import_supports_docx_and_kb_binding():
    app, _db = _build_test_app(_local_tmp_dir())
    client = app.test_client()

    create_kb = client.post(
        "/api/knowledge-bases",
        json={"name": "Requirement KB", "description": "requirement docs"},
    )
    kb_id = create_kb.get_json()["knowledge_base"]["id"]

    response = client.post(
        "/api/requirements/import",
        data={
            "title": "Imported docx requirement",
            "knowledge_base_id": str(kb_id),
            "file": (io.BytesIO(_build_docx_bytes()), "login.docx"),
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 201
    payload = response.get_json()
    assert payload["requirement"]["knowledge_base_id"] == kb_id
    assert "SMS verification" in payload["requirement"]["description"]


def test_knowledge_base_file_import_creates_entry():
    app, _db = _build_test_app(_local_tmp_dir())
    client = app.test_client()

    create_kb = client.post(
        "/api/knowledge-bases",
        json={"name": "Login KB", "description": "login documents"},
    )
    kb_id = create_kb.get_json()["knowledge_base"]["id"]

    response = client.post(
        f"/api/knowledge-bases/{kb_id}/import",
        data={
            "title": "Login spreadsheet",
            "tags": "login,sms,retry",
            "file": (io.BytesIO(_build_xlsx_bytes()), "login.xlsx"),
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 201
    payload = response.get_json()
    assert payload["entry"]["knowledge_base_id"] == kb_id
    assert "Retry limit" in payload["entry"]["content"]


def test_flow_start_supports_knowledge_base_binding(monkeypatch):
    import importlib
    import sys

    import config as config_module

    flow_db_path = _local_tmp_dir() / "phase6_flow_test.db"
    monkeypatch.setenv("DATABASE_URI", f"sqlite:///{flow_db_path.as_posix()}")
    importlib.reload(config_module)
    sys.modules.pop("main", None)
    import main

    client = main.app.test_client()

    with main.app.app_context():
        from models import KnowledgeBase, db

        knowledge_base = KnowledgeBase(name=f"Bound KB {flow_db_path.name}", description="bound")
        db.session.add(knowledge_base)
        db.session.commit()
        knowledge_base_id = knowledge_base.id

    response = client.post(
        "/api/flow/start",
        json={
            "title": "Need login flow",
            "demand": "Login should support sms verification",
            "knowledge_base_id": knowledge_base_id,
        },
    )

    assert response.status_code == 202

    with main.app.app_context():
        from models import Requirement

        requirement = db.session.get(Requirement, response.get_json()["requirement_id"])
        assert requirement.knowledge_base_id == knowledge_base_id
