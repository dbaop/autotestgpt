from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader


def parse_uploaded_file(filename: str, content_bytes: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()

    if suffix in {".txt", ".md", ".markdown", ".csv", ".json", ".log"}:
        return _decode_text(content_bytes)

    if suffix == ".docx":
        return _parse_docx(content_bytes)

    if suffix == ".pdf":
        return _parse_pdf(content_bytes)

    if suffix == ".xlsx":
        return _parse_xlsx(content_bytes)

    raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")


def _decode_text(content_bytes: bytes) -> str:
    try:
        return content_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return content_bytes.decode("gbk", errors="ignore").strip()


def _parse_xlsx(content_bytes: bytes) -> str:
    workbook = load_workbook(BytesIO(content_bytes), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"[Sheet] {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines).strip()


def _parse_docx(content_bytes: bytes) -> str:
    document = Document(BytesIO(content_bytes))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if values:
                lines.append(" | ".join(values))

    return "\n".join(lines).strip()


def _parse_pdf(content_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(content_bytes))
    pages: list[str] = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)
    return "\n".join(pages).strip()
